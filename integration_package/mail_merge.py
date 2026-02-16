#!/usr/bin/env python3
"""
Mail Merge System for Caravan Park Outreach
Generates personalized emails and documents for contacting park operators
"""

import pandas as pd
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from jinja2 import Template
import os
from datetime import datetime
from typing import List, Dict
import json

class MailMergeGenerator:
    """Generate personalized outreach materials for caravan parks"""
    
    def __init__(self, template_dir: str = "templates"):
        """Initialize with template directory"""
        self.template_dir = template_dir
        os.makedirs(template_dir, exist_ok=True)
        self.create_default_templates()
    
    def create_default_templates(self):
        """Create default email and letter templates"""
        
        # Email template for initial contact
        email_template = """
Subject: Partnership Opportunity - {{ park_name }}

Dear {{ park_name }} Management Team,

I hope this message finds you well. My name is {{ sender_name }} from {{ company_name }}, and I'm reaching out regarding exciting development opportunities for caravan and holiday parks in {{ state }}.

{% if development_score > 70 %}
Your property particularly caught our attention due to its excellent development potential, with {{ size_ha|round(1) }} hectares of land in a prime location.
{% else %}
We've identified your {{ size_ha|round(1) }} hectare property as having significant potential for enhancement and modernization.
{% endif %}

We specialize in:
• Modernizing caravan park facilities to meet growing tourism demand
• Developing eco-friendly accommodation options
• Creating revenue-generating amenities
• Improving operational efficiency through smart technology

{% if rating and rating < 4.0 %}
We've noticed there may be opportunities to enhance guest satisfaction and increase your property's market position. Our team has successfully transformed similar properties, achieving average rating improvements of 1.5 stars within 18 months.
{% endif %}

{% if permanently_closed %}
We understand the property may currently be closed. This presents a unique opportunity for comprehensive redevelopment without disrupting existing operations.
{% endif %}

I would love to schedule a brief 15-minute call to discuss how we might work together to unlock your property's full potential. Are you available for a conversation next week?

Best regards,

{{ sender_name }}
{{ sender_title }}
{{ company_name }}
📞 {{ sender_phone }}
📧 {{ sender_email }}

P.S. We're currently offering free feasibility assessments for qualifying properties. This includes market analysis, development recommendations, and ROI projections.
"""
        
        # Formal letter template
        letter_template = """
{{ company_letterhead }}

{{ current_date }}

{{ park_name }}
{{ park_address }}

Dear Sir/Madam,

RE: DEVELOPMENT PARTNERSHIP OPPORTUNITY - {{ park_name|upper }}

I am writing to introduce {{ company_name }} and explore potential partnership opportunities for the development and enhancement of your caravan park property.

ABOUT YOUR PROPERTY
Our research indicates that {{ park_name }} comprises approximately {{ size_ha|round(1) }} hectares in {{ state }}, presenting substantial development potential. {% if land_parcel_ids %}The property encompasses land parcels: {{ land_parcel_ids }}.{% endif %}

OUR PROPOSITION
{{ company_name }} specializes in caravan park development and modernization, with a proven track record of successful projects across Australia. We offer:

1. COMPREHENSIVE DEVELOPMENT SERVICES
   • Master planning and design
   • Infrastructure upgrades
   • Accommodation diversification
   • Amenity enhancement

2. FLEXIBLE PARTNERSHIP MODELS
   • Joint venture developments
   • Management agreements
   • Outright acquisition
   • Lease-to-develop arrangements

3. PROVEN RESULTS
   • Average 40% increase in revenue within 2 years
   • 95% occupancy rates for developed properties
   • Award-winning sustainable design practices

{% if size_ha > 20 %}
LARGE-SCALE DEVELOPMENT OPPORTUNITY
Given your property's substantial size, we envision opportunities for:
• Mixed accommodation offerings (cabins, glamping, RV sites)
• Recreation facilities (pools, playgrounds, activities)
• Commercial amenities (shops, restaurants)
• Potential subdivision for residential development (subject to planning)
{% endif %}

NEXT STEPS
We would welcome the opportunity to:
1. Conduct a complimentary site assessment
2. Present our development concepts
3. Discuss partnership structures that align with your objectives

Please contact me at your earliest convenience to arrange a meeting. I am available to visit your property at a time that suits you.

Thank you for considering this opportunity. I look forward to your response.

Yours sincerely,


{{ sender_name }}
{{ sender_title }}
{{ company_name }}

Direct: {{ sender_phone }}
Email: {{ sender_email }}
"""
        
        # Save templates
        with open(os.path.join(self.template_dir, "email_template.txt"), "w") as f:
            f.write(email_template)
        
        with open(os.path.join(self.template_dir, "letter_template.txt"), "w") as f:
            f.write(letter_template)
    
    def load_template(self, template_name: str) -> str:
        """Load a template from file"""
        template_path = os.path.join(self.template_dir, template_name)
        with open(template_path, "r") as f:
            return f.read()
    
    def generate_email(self, park_data: Dict, sender_info: Dict, template_name: str = "email_template.txt") -> str:
        """Generate personalized email content"""
        template_str = self.load_template(template_name)
        template = Template(template_str)
        
        # Merge data
        merged_data = {**park_data, **sender_info}
        merged_data['current_date'] = datetime.now().strftime("%B %d, %Y")
        
        return template.render(**merged_data)
    
    def generate_word_letter(self, park_data: Dict, sender_info: Dict, output_path: str):
        """Generate a formal Word document letter"""
        doc = Document()
        
        # Set margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        
        # Add letterhead (placeholder)
        header = doc.add_paragraph()
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        header_run = header.add_run(sender_info.get('company_name', 'Company Name'))
        header_run.font.size = Pt(16)
        header_run.font.bold = True
        
        # Add date
        doc.add_paragraph(datetime.now().strftime("%B %d, %Y"))
        
        # Add recipient address
        doc.add_paragraph(park_data.get('park_name', 'Park Name'))
        if park_data.get('park_address'):
            doc.add_paragraph(park_data['park_address'])
        
        doc.add_paragraph()  # Blank line
        
        # Add salutation
        doc.add_paragraph("Dear Sir/Madam,")
        doc.add_paragraph()
        
        # Add subject line
        subject = doc.add_paragraph()
        subject_run = subject.add_run(f"RE: DEVELOPMENT OPPORTUNITY - {park_data.get('park_name', 'Your Park').upper()}")
        subject_run.font.bold = True
        doc.add_paragraph()
        
        # Add body content
        template_str = self.load_template("letter_template.txt")
        template = Template(template_str)
        merged_data = {**park_data, **sender_info}
        merged_data['current_date'] = datetime.now().strftime("%B %d, %Y")
        
        letter_content = template.render(**merged_data)
        
        # Parse and add content (simplified - you might want to handle formatting better)
        for line in letter_content.split('\n'):
            if line.strip():
                doc.add_paragraph(line)
        
        # Save document
        doc.save(output_path)
    
    def generate_bulk_emails(self, parks_df: pd.DataFrame, sender_info: Dict, output_dir: str = "outreach"):
        """Generate emails for multiple parks"""
        os.makedirs(output_dir, exist_ok=True)
        
        generated = []
        
        for idx, park in parks_df.iterrows():
            # Prepare park data
            park_data = {
                'park_name': park.get('Name', 'Caravan Park'),
                'state': park.get('state', 'Australia'),
                'size_ha': park.get('size_ha', park.get('land_area_sqm', 0) / 10000),
                'development_score': park.get('development_score', 50),
                'rating': park.get('rating'),
                'permanently_closed': park.get('permanently_closed', False),
                'land_parcel_ids': park.get('land_parcel_ids', ''),
                'park_address': park.get('formatted_address', '')
            }
            
            # Generate email
            email_content = self.generate_email(park_data, sender_info)
            
            # Save email
            safe_name = "".join(c for c in park_data['park_name'] if c.isalnum() or c in (' ', '-', '_')).rstrip()
            email_file = os.path.join(output_dir, f"email_{safe_name}.txt")
            
            with open(email_file, 'w') as f:
                f.write(email_content)
            
            # Generate Word letter for high-priority parks
            if park_data['development_score'] > 70:
                letter_file = os.path.join(output_dir, f"letter_{safe_name}.docx")
                self.generate_word_letter(park_data, sender_info, letter_file)
            
            generated.append({
                'park_name': park_data['park_name'],
                'email_file': email_file,
                'has_letter': park_data['development_score'] > 70
            })
        
        return generated


def create_campaign_tracker(parks_df: pd.DataFrame, output_file: str = "campaign_tracker.xlsx"):
    """Create an Excel file to track outreach campaign progress"""
    
    # Prepare tracking dataframe
    tracker_df = parks_df[['Name', 'state', 'size_ha', 'phone', 'email', 'website', 'development_score']].copy()
    
    # Add tracking columns
    tracker_df['contact_status'] = 'Not contacted'
    tracker_df['contact_date'] = ''
    tracker_df['contact_method'] = ''
    tracker_df['response_received'] = False
    tracker_df['response_date'] = ''
    tracker_df['response_sentiment'] = ''
    tracker_df['follow_up_required'] = True
    tracker_df['follow_up_date'] = ''
    tracker_df['notes'] = ''
    tracker_df['opportunity_stage'] = 'Initial'
    
    # Sort by development score
    tracker_df = tracker_df.sort_values('development_score', ascending=False)
    
    # Save to Excel with formatting
    with pd.ExcelWriter(output_file, engine='openpyxl') as writer:
        tracker_df.to_excel(writer, sheet_name='Campaign Tracker', index=False)
        
        # Get the workbook and worksheet
        workbook = writer.book
        worksheet = writer.sheets['Campaign Tracker']
        
        # Adjust column widths
        for column in worksheet.columns:
            max_length = 0
            column = [cell for cell in column]
            for cell in column:
                try:
                    if len(str(cell.value)) > max_length:
                        max_length = len(str(cell.value))
                except:
                    pass
            adjusted_width = min(max_length + 2, 50)
            worksheet.column_dimensions[column[0].column_letter].width = adjusted_width
    
    return tracker_df


def main():
    """Main execution for mail merge generation"""
    
    print("📧 Caravan Parks Mail Merge Generator")
    print("=" * 50)
    
    # Load the enriched data
    files = [f for f in os.listdir('.') if f.startswith('enriched_caravan_parks_') and f.endswith('.xlsx')]
    
    if not files:
        print("❌ No enriched data file found. Please run enrich_with_google.py first.")
        return
    
    latest_file = sorted(files)[-1]
    print(f"Loading data from: {latest_file}")
    df = pd.read_excel(latest_file)
    
    # Calculate size in hectares if not present
    if 'size_ha' not in df.columns:
        df['size_ha'] = df['land_area_sqm'] / 10000
    
    # Get sender information
    print("\n📝 Enter your details for the mail merge:")
    sender_info = {
        'sender_name': input("Your name: ") or "John Smith",
        'sender_title': input("Your title: ") or "Development Director",
        'company_name': input("Company name: ") or "Park Developments Australia",
        'sender_phone': input("Your phone: ") or "0400 000 000",
        'sender_email': input("Your email: ") or "john@parkdev.com.au",
        'company_letterhead': input("Company address (for letters): ") or "123 Business St, Sydney NSW 2000"
    }
    
    # Select parks for outreach
    print("\n🎯 Select parks for outreach:")
    print("1. Top 20 by development score")
    print("2. All parks with contact details")
    print("3. Custom filter")
    
    choice = input("Enter choice (1-3): ") or "1"
    
    if choice == "1":
        # Top 20 by score
        if 'development_score' in df.columns:
            outreach_df = df.nlargest(20, 'development_score')
        else:
            outreach_df = df.nlargest(20, 'size_ha')
    elif choice == "2":
        # Parks with contact details
        outreach_df = df[(df['phone'].notna()) | (df['email'].notna()) | (df['website'].notna())]
    else:
        # Custom filter
        min_size = float(input("Minimum size (hectares): ") or "20")
        outreach_df = df[df['size_ha'] >= min_size]
    
    print(f"\n✅ Selected {len(outreach_df)} parks for outreach")
    
    # Generate mail merge materials
    generator = MailMergeGenerator()
    
    print("\n📄 Generating personalized emails and letters...")
    generated = generator.generate_bulk_emails(outreach_df, sender_info)
    
    print(f"✅ Generated {len(generated)} email templates")
    print(f"✅ Generated {sum(1 for g in generated if g['has_letter'])} formal letters (high-priority parks)")
    
    # Create campaign tracker
    print("\n📊 Creating campaign tracker...")
    tracker = create_campaign_tracker(outreach_df)
    print("✅ Campaign tracker saved to campaign_tracker.xlsx")
    
    print("\n🎉 Mail merge complete!")
    print("\nNext steps:")
    print("1. Review generated emails in the 'outreach' folder")
    print("2. Use campaign_tracker.xlsx to track your outreach")
    print("3. Send emails using your preferred email client")
    print("4. Print letters for high-priority prospects")
    
    # Save sender info for future use
    with open('sender_info.json', 'w') as f:
        json.dump(sender_info, f)
    print("\n💾 Your details saved for future use")


if __name__ == "__main__":
    main()
